from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DESKTOP = Path(r"C:\Users\tbond\OneDrive - Fundación UADE\Escritorio")
OUTPUT = DESKTOP / "TPO_G1_MuJoCo_Guia_Docente.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(90, 90, 90)
FILL_BLUE = "E8EEF5"
FILL_GRAY = "F2F4F7"
FILL_NOTE = "F7F9FC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_indent(table, dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run(paragraph, text: str, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_paragraph(doc, text="", style=None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def add_bullet(doc, text: str, level: int = 0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    return add_paragraph(doc, text, style=style)


def add_number(doc, text: str):
    return add_paragraph(doc, text, style="List Number")


def add_code_block(doc, code: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_indent(table)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.splitlines()):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
    return table


def add_note_box(doc, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_indent(table)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, FILL_NOTE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title, bold=True, color=DARK_BLUE)
    p.add_run(" ")
    p.add_run(body)
    return table


def add_label_detail_table(doc, rows: list[tuple[str, str]], header: tuple[str, str] | None = None):
    table = doc.add_table(rows=1 if header else 0, cols=2)
    table.style = "Table Grid"
    set_table_indent(table)
    set_table_width(table, [1.85, 4.65])
    if header:
        row = table.rows[0]
        set_repeat_table_header(row)
        for idx, text in enumerate(header):
            set_cell_shading(row.cells[idx], FILL_BLUE)
            row.cells[idx].paragraphs[0].add_run(text).bold = True
    for label, detail in rows:
        row = table.add_row()
        set_cell_shading(row.cells[0], FILL_GRAY)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(detail)
    set_table_width(table, [1.85, 4.65])
    return table


def add_three_col_table(doc, headers: tuple[str, str, str], rows: list[tuple[str, str, str]], widths=(1.3, 2.45, 2.75)):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_indent(table)
    set_table_width(table, list(widths))
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for idx, header in enumerate(headers):
        set_cell_shading(header_row.cells[idx], FILL_BLUE)
        header_row.cells[idx].paragraphs[0].add_run(header).bold = True
    for values in rows:
        row = table.add_row()
        for idx, value in enumerate(values):
            row.cells[idx].paragraphs[0].add_run(value)
    set_table_width(table, list(widths))
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UADE Robot Lab - Guia docente G1 MuJoCo - Version visual/cinematica")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Bullet 2", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(title, "TPO - Programacion secuencial con G1 en MuJoCo", bold=True, color=DARK_BLUE, size=22)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    add_run(subtitle, "Guia docente para preparar y ejecutar una actividad visual con robot humanoide simulado", italic=True, color=MUTED, size=11)

    add_label_detail_table(
        doc,
        [
            ("Materia", "Fundamentos de Informatica / Introduccion a la Programacion"),
            ("Duracion sugerida", "2 horas de taller presencial + extension opcional como TPO"),
            ("Robot", "Unitree G1 en MuJoCo, modo visual/cinematico docente"),
            ("Modalidad", "Individual o parejas, con una PC del profesor proyectando el simulador"),
            ("Nivel", "Inicial: programacion secuencial, funciones y parametros"),
        ],
    )

    add_note_box(
        doc,
        "Estado de esta version:",
        "el robot se controla de manera visual/cinematica. El comando movimiento traslada la base y anima articulaciones para que los alumnos vean consecuencias claras de sus instrucciones. No es locomocion fisica realista ni sim-to-real.",
    )

    add_heading(doc, "1. Objetivo de la actividad", 1)
    add_paragraph(
        doc,
        "El objetivo es que los alumnos escriban una secuencia simple de instrucciones en Python y observen el efecto visual sobre un robot humanoide G1 simulado en MuJoCo.",
    )
    for item in [
        "Comprender que un programa ejecuta instrucciones en un orden determinado.",
        "Usar funciones de alto nivel con parametros simples.",
        "Relacionar codigo Python con acciones visibles del robot.",
        "Practicar prueba, observacion y correccion a partir de logs y comportamiento visual.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Archivos incluidos", 1)
    add_label_detail_table(
        doc,
        [
            ("setup_windows.bat", "Instalador para PC Windows. Crea .venv, instala dependencias y verifica el simulador."),
            ("abrir_g1_sim.bat", "Abre el simulador con doble clic o desde consola."),
            ("run_g1_sim.ps1", "Launcher PowerShell del simulador G1."),
            ("g1_teacher_sim.py", "Runner docente que abre MuJoCo y expone la API local."),
            ("g1_student_api.py", "API simple que importan los alumnos."),
            ("examples/ejemplo_g1_simple.py", "Ejemplo listo para probar todos los endpoints principales."),
            ("examples/alumnos_g1_api.py", "Plantilla editable para alumnos."),
        ],
        header=("Archivo", "Uso"),
    )

    add_heading(doc, "3. Requisitos para la PC del profesor", 1)
    for item in [
        "Windows 10/11.",
        "Python 3.10 instalado y disponible como py -3.10.",
        "Git para Windows instalado y disponible en PATH.",
        "Conexion a internet para la primera instalacion.",
        "Repositorio UadeRobotLab descargado o clonado.",
    ]:
        add_bullet(doc, item)
    add_note_box(
        doc,
        "Importante:",
        "el instalador usa --trusted-host para PyPI porque algunas redes institucionales interceptan certificados SSL y pip puede fallar si no se contempla ese caso.",
    )

    add_heading(doc, "4. Instalacion inicial", 1)
    add_paragraph(doc, "Estos pasos se hacen una sola vez por PC.")
    add_number(doc, "Abrir PowerShell o CMD en la carpeta del repositorio.")
    add_number(doc, "Entrar a la carpeta del simulador.")
    add_code_block(
        doc,
        r"""cd 04Simuladores\UnitreeMujocoOficial
setup_windows.bat""",
    )
    add_number(doc, "Esperar a que termine. El instalador debe finalizar con Setup terminado.")
    add_number(doc, "Si falla por Python o Git, instalar el requisito indicado y volver a ejecutar.")

    add_heading(doc, "5. Abrir el simulador para la clase", 1)
    add_paragraph(doc, "Antes de que los alumnos ejecuten scripts, el simulador debe quedar abierto.")
    add_code_block(
        doc,
        r"""cd 04Simuladores\UnitreeMujocoOficial
abrir_g1_sim.bat""",
    )
    add_paragraph(doc, "Alternativa desde PowerShell:")
    add_code_block(doc, r""".\run_g1_sim.ps1""")
    add_note_box(
        doc,
        "Ventana abierta:",
        "mientras la ventana de MuJoCo este abierta, la API local queda escuchando comandos en 127.0.0.1:8765.",
    )

    add_heading(doc, "6. Probar el ejemplo simple", 1)
    add_paragraph(doc, "Con MuJoCo abierto, en otra terminal ejecutar:")
    add_code_block(doc, r""".\.venv\Scripts\python.exe .\examples\ejemplo_g1_simple.py""")
    add_paragraph(doc, "El ejemplo realiza la siguiente secuencia:")
    for item in ["conectar", "girar", "saludar", "caminar hacia adelante", "dar un beso", "detenerse", "desconectar"]:
        add_bullet(doc, item)

    add_heading(doc, "7. API disponible para alumnos", 1)
    add_label_detail_table(
        doc,
        [
            ("conectar()", "Conecta el script del alumno con el simulador abierto."),
            ("verificar_estado()", "Devuelve posicion y accion actual del robot."),
            ("movimiento(adelante, costado, giro, tiempo)", "Mueve visualmente al robot durante una cantidad de segundos."),
            ("saludar()", "Anima el brazo derecho para saludar."),
            ("dar_beso()", "Anima el brazo derecho para simular un beso."),
            ("detenerse()", "Frena el movimiento y vuelve a pose estable."),
            ("desconectar()", "Cierra logicamente la conexion del script."),
        ],
        header=("Metodo", "Descripcion"),
    )
    add_paragraph(doc, "Parametros recomendados para movimiento:")
    add_three_col_table(
        doc,
        ("Parametro", "Rango sugerido", "Ejemplo"),
        [
            ("adelante", "-0.45 a 0.45", "0.20 para avanzar despacio"),
            ("costado", "-0.30 a 0.30", "0.15 para moverse lateralmente"),
            ("giro", "-1.00 a 1.00", "0.50 para girar a la izquierda"),
            ("tiempo", "0 a 10 segundos", "3.0 para ejecutar durante 3 segundos"),
        ],
        widths=(1.45, 2.05, 3.0),
    )

    add_heading(doc, "8. Ejemplo de codigo para alumnos", 1)
    add_paragraph(doc, "Este ejemplo es intencionalmente corto para que se vea la programacion secuencial.")
    add_code_block(
        doc,
        """from g1_student_api import RobotG1


robot = RobotG1()

robot.conectar()

print("El robot va a girar un poco.")
robot.movimiento(adelante=0.0, costado=0.0, giro=0.5, tiempo=4.0)

print("El robot va a saludar.")
robot.saludar()

print("El robot va a caminar hacia adelante.")
robot.movimiento(adelante=0.2, costado=0.0, giro=0.0, tiempo=3.0)

print("El robot va a dar un beso.")
robot.dar_beso()

print("El robot se detiene.")
robot.detenerse()

robot.desconectar()
""",
    )

    add_heading(doc, "9. Desarrollo sugerido del taller", 1)
    add_three_col_table(
        doc,
        ("Tiempo", "Docente", "Alumnos"),
        [
            ("0-10 min", "Presenta el objetivo y muestra el G1 en MuJoCo.", "Observan el simulador y preguntan."),
            ("10-20 min", "Ejecuta ejemplo_g1_simple.py y relaciona codigo con accion.", "Identifican orden de instrucciones."),
            ("20-35 min", "Explica conectar, movimiento, saludar, detenerse.", "Preparan su archivo Python."),
            ("35-75 min", "Guia una secuencia paso a paso.", "Programan giro, saludo y avance."),
            ("75-95 min", "Propone variar parametros.", "Cambian tiempo, giro y velocidad."),
            ("95-110 min", "Agrega desafio opcional: dar_beso o repetir una accion.", "Extienden la rutina."),
            ("110-120 min", "Cierre y reflexion sobre orden y parametros.", "Muestran 2 o 3 rutinas."),
        ],
        widths=(1.05, 2.75, 2.7),
    )

    add_heading(doc, "10. Consigna breve para alumnos", 1)
    add_paragraph(
        doc,
        "Implementa en Python una rutina secuencial para el robot G1 simulado. La rutina debe conectarse al simulador, girar, saludar, avanzar, ejecutar una accion final y detenerse.",
    )
    for item in [
        "Usar al menos cuatro llamadas a metodos del robot.",
        "Incluir mensajes print que expliquen cada paso.",
        "Usar parametros de movimiento seguros.",
        "Entregar el archivo .py y una captura o descripcion de la ejecucion.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Criterios de evaluacion", 1)
    add_three_col_table(
        doc,
        ("Criterio", "Indicador", "Puntaje"),
        [
            ("Secuencia", "El programa ejecuta acciones en orden claro y sin errores.", "35"),
            ("Uso de API", "Usa correctamente conectar, movimiento, saludar/dar_beso y detenerse.", "25"),
            ("Parametros", "Elige valores razonables de tiempo, giro y avance.", "15"),
            ("Legibilidad", "Codigo simple, nombres claros y mensajes print utiles.", "15"),
            ("Reflexion", "Explica que cambia si altera el orden o los parametros.", "10"),
        ],
        widths=(1.35, 4.25, 0.9),
    )

    add_heading(doc, "12. Problemas frecuentes", 1)
    add_three_col_table(
        doc,
        ("Problema", "Causa probable", "Solucion"),
        [
            ("py no se reconoce", "Python no instalado o no esta en PATH.", "Instalar Python 3.10 y marcar Add to PATH."),
            ("git no se reconoce", "Git no instalado.", "Instalar Git para Windows."),
            ("pip falla por SSL", "Red institucional con certificado interceptado.", "Usar setup_windows.bat; ya incluye trusted-host."),
            ("No se abre MuJoCo", "Dependencia faltante o instalacion incompleta.", "Ejecutar setup_windows.bat y revisar el mensaje de error."),
            ("El ejemplo no conecta", "MuJoCo no esta abierto o el puerto cambio.", "Abrir run_g1_sim.ps1 antes de correr el script."),
            ("Puerto ocupado", "Otro proceso usa 8765.", "Abrir con .\\run_g1_sim.ps1 -ApiPort 8766 y ajustar RobotG1(port=8766)."),
            ("Movimiento raro/flotante", "La version es visual/cinematica.", "Para este TPO es esperado; no representa fisica realista."),
        ],
        widths=(1.65, 2.25, 2.6),
    )

    add_heading(doc, "13. Notas pedagogicas", 1)
    for item in [
        "Enfatizar que el foco es el orden de instrucciones, no la fisica del robot.",
        "Pedir que los alumnos modifiquen un parametro por vez y observen el resultado.",
        "Usar errores como parte de la clase: simulador cerrado, metodo mal escrito, parametros fuera de rango.",
        "Aclarar que una caminata fisica real requiere controladores de locomocion o politicas entrenadas.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "14. Comandos rapidos", 1)
    add_code_block(
        doc,
        r""":: Instalacion inicial
cd 04Simuladores\UnitreeMujocoOficial
setup_windows.bat

:: Abrir simulador
abrir_g1_sim.bat

:: Verificar sin abrir MuJoCo
powershell -NoProfile -ExecutionPolicy Bypass -File .\run_g1_sim.ps1 -SetupOnly

:: Probar ejemplo con MuJoCo abierto
.\.venv\Scripts\python.exe .\examples\ejemplo_g1_simple.py
""",
    )

    add_note_box(
        doc,
        "Cierre:",
        "programar un robot en esta actividad significa ordenar instrucciones, elegir parametros y observar consecuencias. Aunque el robot sea visual, el razonamiento secuencial es el mismo.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
